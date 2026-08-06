"""
doc_parser.py

Parses the Work Instruction .docx files into structured fields.

The Work Instructions have a fairly consistent structure:
  - "Operations/Work/Job Activity covered by this assessment: <TITLE>"
  - "Scope:"
  - "Required Personal Protective Equipment (PPE):"
  - "Loading & Unloading Check:"
  - "Operations Procedure:"
  - "Shutdown:"
  - "Inspection:"
  - Assessor / Signature / Approved blocks

We extract the raw text with python-docx, then split into sections and
build a structured WorkInstruction record.
"""

import os
import re
from typing import Dict, List, Optional

from docx import Document


HEADING_PATTERNS = [
    r"Required\s*Personal\s*Protective\s*Equipment\s*\(?\s*PPE\s*\)?\s*:",
    r"Loading\s*&\s*Unloading\s*Check\s*:",
    r"Operations?\s*Procedure\s*:",
    r"Operation\s*Procedure\s*:",
    r"Shutdown\s*:",
    r"Inspection\s*:",
    r"Scope\s*:",
    r"Prerequisites?\s*:",
    r"Pre[- ]?start\s*Checks?\s*:",
    r"Machine\s*Parameters?\s*:",
    r"Quality\s*Requirements?\s*:",
    r"Acceptance\s*Criteria\s*:",
    r"Safety\s*Notes?\s*:",
    r"Tools?\s*Required\s*:",
    r"Consumables?\s*:",
    r"Applicable\s*Machines?\s*:",
    r"Revision\s*History\s*:",
]

SECTION_KEYS = [
    "scope",
    "ppe",
    "pre_start_checks",
    "prerequisites",
    "machine_parameters",
    "tools_required",
    "consumables",
    "applicability",
    "procedure",
    "inspection",
    "quality_requirements",
    "acceptance_criteria",
    "shutdown_procedure",
    "safety_notes",
    "revision_history",
]


def extract_text_from_docx(path: str) -> str:
    """Extract all paragraph text from a .docx file."""
    doc = Document(path)
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)
    # Also extract table content
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append(" | ".join([c for c in cells if c]))
    return "\n".join(lines)


def parse_wi_number(filename: str) -> str:
    """Extract the WI number from a filename like 'WI_04 for Blasting.docx'."""
    match = re.search(r"(WI[_\- ]?\d+)", filename, re.IGNORECASE)
    if match:
        return match.group(1).replace("_", "").replace("-", "").replace(" ", "")
    return "WI"


def normalize_title_fragment(title_fragment: str) -> str:
    """Normalize repeated title header fragments, Work Instruction prefixes, and pipe-delimited duplicates."""
    title_fragment = title_fragment.strip()
    title_fragment = re.sub(
        r"(?:Work\s*Instruction\s*(?:for|-|:)?\s*)+",
        "",
        title_fragment,
        flags=re.IGNORECASE,
    ).strip()
    title_fragment = re.sub(
        r"(?:Operations?/Work/Job\s*Activity\s*covered\s*by\s*this\s*assessment\s*:\s*)+",
        "",
        title_fragment,
        flags=re.IGNORECASE,
    ).strip()

    parts = [part.strip() for part in re.split(r"\s*\|\s*", title_fragment) if part.strip()]
    seen = []
    for part in parts:
        if part.lower() not in [s.lower() for s in seen]:
            seen.append(part)
    res = seen[0] if seen else title_fragment
    res = re.sub(r"^(?:Work\s*Instruction\s*(?:for|-|:)?\s*)+", "", res, flags=re.IGNORECASE).strip()
    return res


def _clean_raw_text(raw_text: str) -> str:
    """Clean repeated header fragments and duplicate lines from raw extracted text."""
    cleaned = re.sub(
        r"(?:Operations?/Work/Job\s*Activity\s*covered\s*by\s*this\s*assessment\s*:\s*)+",
        "",
        raw_text,
        flags=re.IGNORECASE,
    )
    cleaned = re.sub(r"\s*\|\s*", " | ", cleaned)
    cleaned_lines = []
    seen = set()
    for line in cleaned.splitlines():
        normalized = re.sub(r"\s+", " ", line).strip().lower()
        if normalized and normalized not in seen:
            seen.add(normalized)
            cleaned_lines.append(line)
    return "\n".join(cleaned_lines).strip()


def parse_title(raw_text: str, filename: str) -> str:
    """Extract the work/job activity title."""
    # Look for the standard header line
    match = re.search(
        r"Operations?/Work/Job\s*Activity\s*covered\s*by\s*this\s*assessment\s*:\s*([^\n\r]+)",
        raw_text,
        re.IGNORECASE,
    )
    if match:
        return normalize_title_fragment(match.group(1))
    # Fallback: derive from filename
    name = filename
    name = re.sub(r"\.docx?$", "", name, flags=re.IGNORECASE)
    name = re.sub(r"^WI_?\d+\s*for\s*", "", name, flags=re.IGNORECASE)
    name = re.sub(r"^WI_?\d+\s*", "", name, flags=re.IGNORECASE)
    name = name.replace("_", " ").strip()
    return name or "Unknown Work Instruction"


def is_heading_line(line: str) -> Optional[str]:
    """Return the section key if the line matches a heading pattern, else None."""
    for pattern in SECTION_KEYS:
        # Build a regex from the key
        label = pattern.replace("_", " ")
        if label == "ppe":
            label = "Required Personal Protective Equipment (PPE)"
        elif label == "pre start checks":
            label = "Pre-start Checks"
        elif label == "machines":
            label = "Applicable Machines"
        elif label == "tools required":
            label = "Tools Required"
        elif label == "consumables":
            label = "Consumables"
        elif label == "quality requirements":
            label = "Quality Requirements"
        elif label == "acceptance criteria":
            label = "Acceptance Criteria"
        elif label == "shutdown procedure":
            label = "Shutdown"
        elif label == "safety notes":
            label = "Safety Notes"
        elif label == "revision history":
            label = "Revision History"
        elif label == "prerequisites":
            label = "Prerequisites"
        elif label == "inspection":
            label = "Inspection"
        elif label == "scope":
            label = "Scope"
        elif label == "procedure":
            label = "Operations Procedure"

        if re.search(label, line, re.IGNORECASE):
            return pattern
    return None


def determine_department(filename: str) -> str:
    """Infer department from filename keywords."""
    lower = filename.lower()
    if any(k in lower for k in ["spray", "blasting", "grinding", "machining", "cutting", "polishing", "mounting"]):
        return "Spray / Surface Engineering"
    if any(k in lower for k in ["inward", "outward", "dispatch", "packing", "handling"]):
        return "Logistics / Stores"
    if any(k in lower for k in ["inspection", "calibration", "test", "visual"]):
        return "Quality"
    if any(k in lower for k in ["ppe", "chemical", "cleaning"]):
        return "Safety / EHS"
    return "Production"


def parse_work_instruction(docx_path: str) -> Dict:
    """Parse a single .docx work instruction into a structured dict."""
    raw_text = extract_text_from_docx(docx_path)
    raw_text = _clean_raw_text(raw_text)
    filename = os.path.basename(docx_path)

    result = {
        "wi_number": parse_wi_number(filename),
        "title": parse_title(raw_text, filename),
        "department": determine_department(filename),
        "file_path": docx_path,
        "raw_text": raw_text,
        "sections": [],
    }

    # Try to find a revision indicator
    rev_match = re.search(r"[Rr]ev(?:ision)?\.?\s*[:#]?\s*([\w.\-]+)", raw_text)
    if rev_match:
        result["revision"] = f"Rev {rev_match.group(1)}"
    else:
        result["revision"] = "Rev 1"

    # Split into sections by headings
    lines = raw_text.split("\n")
    current_key = None
    current_lines = []

    section_order = 0
    for line in lines:
        key = is_heading_line(line)
        if key:
            # save previous
            if current_key and current_lines:
                content = "\n".join(current_lines).strip()
                if content:
                    result["sections"].append({
                        "heading": current_key,
                        "content": content,
                        "order_index": section_order,
                    })
                    section_order += 1
            current_key = key
            current_lines = []
        else:
            current_lines.append(line)

    # save last section
    if current_key and current_lines:
        content = "\n".join(current_lines).strip()
        if content:
            result["sections"].append({
                "heading": current_key,
                "content": content,
                "order_index": section_order,
            })

    # Map sections to structured fields
    section_map = {s["heading"]: s["content"] for s in result["sections"]}
    for key in SECTION_KEYS:
        if key in section_map:
            result[key] = section_map[key]

    # If no separate scope section, put a concise summary without repeated wording.
    if not result.get("scope"):
        result["scope"] = f"{result['title']} approved procedure and requirements."

    # Determine approval requirements
    result["supervisor_approval_required"] = "Approved" in raw_text or "approval" in raw_text.lower()
    result["qa_approval_required"] = "QA" in raw_text or "Quality" in raw_text

    return result


def parse_all_work_instructions(directory: str) -> List[Dict]:
    """Parse all .docx work instructions in a directory."""
    parsed = []
    if not os.path.isdir(directory):
        return parsed
    for fname in sorted(os.listdir(directory)):
        if fname.lower().endswith((".docx", ".doc")):
            path = os.path.join(directory, fname)
            try:
                parsed.append(parse_work_instruction(path))
            except Exception as e:
                print(f"[doc_parser] Failed to parse {fname}: {e}")
    return parsed
